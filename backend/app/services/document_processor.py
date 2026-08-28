import io
import json
import re
from pathlib import Path


import fitz
import pytesseract 
from PIL import Image
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.redis import redis_client

MIN_EXTRACTABLE_TEXT_LENGTH = 50


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Collapse repeated spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)

    # Avoid excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def extract_page_with_ocr(
    page: fitz.Page,
) -> str:
    pixmap = page.get_pixmap(
        matrix=fitz.Matrix(2, 2),
        alpha=False,
    )

    image = Image.open(
        io.BytesIO(pixmap.tobytes("png"))
    )

    text = pytesseract.image_to_string(
        image,
        lang="eng",
    )

    return clean_text(text)


# Pdf extraction

# def extract_pdf(
#     file_path: str,
#     document_id: str,
# ) -> list[Document]:
#     documents = []

#     pdf = fitz.open(file_path)

#     try:
#         for page_index, page in enumerate(pdf):
#             text = page.get_text("text")
#             text = clean_text(text)

#             if not text:
#                 # OCR fallback will be added shortly.
#                 continue

#             documents.append(
#                 Document(
#                     page_content=text,
#                     metadata={
#                         "document_id": document_id,
#                         "file_type": "pdf",
#                         "page_number": page_index + 1,
#                     },
#                 )
#             )

#     finally:
#         pdf.close()

#     return documents


def extract_pdf(
    file_path: str,
    document_id: str,
) -> list[Document]:
    documents = []

    pdf = fitz.open(file_path)

    try:
        for page_index, page in enumerate(pdf):
            text = page.get_text("text")
            text = clean_text(text)

            extraction_method = "native"

            # OCR fallback for scanned/sparse pages
            if len(text) < MIN_EXTRACTABLE_TEXT_LENGTH:
                try:
                    ocr_text = extract_page_with_ocr(page)

                    if len(ocr_text) > len(text):
                        text = ocr_text
                        extraction_method = "ocr"

                except Exception as exc:
                    print(
                        f"OCR failed on page "
                        f"{page_index + 1}: {exc}"
                    )

            if not text:
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "document_id": document_id,
                        "file_type": "pdf",
                        "page_number": page_index + 1,
                        "extraction_method": extraction_method,
                    },
                )
            )

    finally:
        pdf.close()

    return documents




# Docx extraction

def extract_docx(
    file_path: str,
    document_id: str,
) -> list[Document]:
    docx = DocxDocument(file_path)

    documents = []

    current_heading = None
    current_content = []
    section_index = 1

    def save_section():
        nonlocal section_index
        nonlocal current_content

        text = clean_text("\n".join(current_content))

        if not text:
            return

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "document_id": document_id,
                    "file_type": "docx",
                    "section_index": section_index,
                    "section_heading": current_heading,
                },
            )
        )

        section_index += 1
        current_content = []

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        style_name = (
            paragraph.style.name
            if paragraph.style
            else ""
        )

        if style_name.lower().startswith("heading"):
            if current_content:
                save_section()

            current_heading = text

        else:
            current_content.append(text)

    
    for table in docx.tables:
        table_rows = []

        for row in table.rows:
            cells = [
                clean_text(cell.text)
                for cell in row.cells
            ]

            table_rows.append(" | ".join(cells))

        table_text = "\n".join(table_rows)

        if table_text.strip():
            current_content.append(table_text)

    if current_content:
        save_section()

    return documents



# Create the format router

def extract_document(
    file_path: str,
    file_type: str,
    document_id: str,
) -> list[Document]:
    if file_type == "pdf":
        return extract_pdf(
            file_path,
            document_id,
        )

    if file_type == "docx":
        return extract_docx(
            file_path,
            document_id,
        )

    raise ValueError(
        f"Unsupported file type: {file_type}"
    )


# Add LangChain chunking

def split_documents(
    documents: list[Document],
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            "",
        ],
    )

    chunks = splitter.split_documents(documents)

    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = (
            f"{chunk.metadata['document_id']}_{index}"
        )

    return chunks



async def update_document_metadata(
    document_id: str,
    **updates,
) -> dict:
    redis_key = f"document:{document_id}"

    raw_metadata = await redis_client.get(redis_key)

    if raw_metadata is None:
        raise ValueError(
            f"Document not found: {document_id}"
        )

    metadata = json.loads(raw_metadata)

    metadata.update(updates)

    await redis_client.set(
        redis_key,
        json.dumps(metadata),
        ex=settings.SESSION_EXPIRY_SECONDS,
    )

    return metadata



# Create the processing pipeline

async def process_document(
    document_id: str,
) -> None:
    try:
        redis_key = f"document:{document_id}"

        raw_metadata = await redis_client.get(
            redis_key
        )

        if raw_metadata is None:
            return

        metadata = json.loads(raw_metadata)

        file_path = metadata["storage_path"]
        file_type = metadata["file_type"]

        # Stage 1
        await update_document_metadata(
            document_id,
            status="processing",
            status_detail="extracting_text",
        )

        documents = extract_document(
            file_path=file_path,
            file_type=file_type,
            document_id=document_id,
        )

        if not documents:
            await update_document_metadata(
                document_id,
                status="failed",
                status_detail="failed",
                error="no_extractable_text",
            )
            return

        # Stage 2
        await update_document_metadata(
            document_id,
            status_detail="splitting_document",
        )

        chunks = split_documents(documents)

        if not chunks:
            await update_document_metadata(
                document_id,
                status="failed",
                status_detail="failed",
                error="no_extractable_text",
            )
            return

        # Step 5 will continue from here.
        await update_document_metadata(
            document_id,
            status="processing",
            status_detail="creating_embeddings",
            extracted_units=len(documents),
            chunk_count=len(chunks),
        )

        print(
            f"[Document {document_id}] "
            f"Extracted {len(documents)} units, "
            f"created {len(chunks)} chunks."
        )

    except Exception as exc:
        print(
            f"Document processing failed: {exc}"
        )

        try:
            await update_document_metadata(
                document_id,
                status="failed",
                status_detail="failed",
                error="document_processing_failed",
            )
        except Exception:
            pass


